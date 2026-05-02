from docx import Document

HEADING_MAP = {
    "Heading 1": "#",
    "Heading 2": "##",
    "Heading 3": "###",
    "Heading 4": "####",
    "Heading 5": "#####",
    "Heading 6": "######",
}


def convert_docx(filepath: str) -> str:
    doc = Document(filepath)
    parts = []
    in_list = False
    list_items = []

    def flush_list():
        nonlocal in_list, list_items
        if list_items:
            parts.append("\n".join(list_items) + "\n")
            list_items.clear()
            in_list = False

    for para in doc.paragraphs:
        style = para.style.name

        if style in HEADING_MAP:
            flush_list()
            prefix = HEADING_MAP[style]
            parts.append(f"\n{prefix} {para.text}\n")
        elif style.startswith("List"):
            in_list = True
            list_items.append(f"- {para.text}")
        elif para.text.strip():
            flush_list()
            if para.runs and para.runs[0].bold and not style.startswith("Heading"):
                parts.append(f"\n**{para.text}**\n")
            else:
                parts.append(f"\n{para.text}\n")

    flush_list()

    for table in doc.tables:
        rows = []
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip().replace("|", "\\|") for cell in row.cells]
            line = "| " + " | ".join(cells) + " |"
            rows.append(line)
            if i == 0:
                rows.append("| " + " | ".join(["---"] * len(cells)) + " |")
        parts.append("\n" + "\n".join(rows) + "\n")

    return "\n".join(parts).strip()
